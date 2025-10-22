from groq import Groq, RateLimitError
from pydantic import BaseModel
import json
import PyPDF2
import os
import time
from dotenv import load_dotenv
from API_key_manager import api_manager, make_api_call_with_retry
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

load_dotenv()

def extract_text_to_json(file_path, json_file):
    """Extract text from PDF or DOCX and save to JSON file"""
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        text = ""
        
        if file_ext == '.pdf':
            # Extract from PDF
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
        
        elif file_ext in ['.docx', '.doc']:
            # Extract from DOCX
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx library not installed. Install it with: pip install python-docx")
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: .pdf, .docx")
        
        if not text.strip():
            raise ValueError(f"No text extracted from {file_path}")
        
        with open(json_file, "w", encoding="utf-8") as jf:
            json.dump({"text": text}, jf, ensure_ascii=False, indent=2)
        return True
        
    except Exception as e:
        raise Exception(f"Error extracting text: {str(e)}")

def Clause_extraction(file):
    """Extract clauses from PDF or DOCX file with automatic API key rotation"""

    class ClauseExtraction(BaseModel):
        clause_id: str
        heading: str
        text: str

    try:
        if not os.path.exists(file):
            raise FileNotFoundError(f"File not found: {file}")

        file_ext = os.path.splitext(file)[1].lower()
        text = ""
        
        if file_ext == '.pdf':
            with open(file, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
        
        elif file_ext in ['.docx', '.doc']:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx library not installed. Install it with: pip install python-docx")
            doc = Document(file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: .pdf, .docx")

        if not text.strip():
            raise ValueError("No text could be extracted from the file")

        prompt = f"""
        You are an expert in legal contract analysis.
        Your task is to extract all **clauses** from the following contract text.
        
        ### Guidelines:
        - A clause may begin with:
          - A number/letter (e.g. "1.", "A."),
          - The word "Clause" followed by a number (e.g. "Clause 1", "Clause 5"), OR
          - An ALL CAPS heading (e.g. "DEFINATION", "TRANSFER OF DATA".)

        - Each extracted clause must include:
          - **clause_id** (the exact numbering/label from the contract)
          - **heading/title** (use the explicit heading if present)
          - **full text** (the complete text of the clause, including sub-clauses)

        - Maintain clause boundaries precisely
        - Include all important clauses including exhibits and appendices
        - Exclude non-contractual content (page numbers, headers, etc.)
        - Response in **valid json** only

        Input: {text}

        Response format:
        [
          {{
            "clause_id": "<clause_id>",
            "heading/title": "<heading>",
            "full text": "<complete_text>"
          }}
        ]
        """

        # Make API call with automatic retry and key rotation
        def api_call(client):
            chat_completion = client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.3-70b-versatile",
                response_format={"type": "json_object"},
                temperature=0.3
            )
            return chat_completion.choices[0].message.content
        
        return make_api_call_with_retry(api_call, max_retries=5)

    except Exception as e:
        raise Exception(f"Error in clause extraction: {str(e)}")

def Clause_extraction_with_summarization(file):
    """Extract and summarize clauses from PDF or DOCX file with automatic API key rotation"""
    try:
        if not os.path.exists(file):
            raise FileNotFoundError(f"File not found: {file}")

        file_ext = os.path.splitext(file)[1].lower()
        text = ""
        
        if file_ext == '.pdf':
            with open(file, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
        
        elif file_ext in ['.docx', '.doc']:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx library not installed. Install it with: pip install python-docx")
            doc = Document(file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: .pdf, .docx")

        if not text.strip():
            raise ValueError("No text could be extracted from the file")

        prompt = f"""
        You are an expert in legal contract analysis.
        Your task is to extract and summarize all clauses from the following contract text.
        
        ### Guidelines:
        - A clause may begin with:
          - A number/letter (e.g. "1.", "A."),
          - The word "Clause" followed by a number (e.g. "Clause 1", "Clause 5"), OR
          - An ALL CAPS heading (e.g. "DEFINATION", "TRANSFER OF DATA".)

        - Each extracted clause must include:
          - **clause_id** (exact numbering/label)
          - **heading/title** (explicit heading or first few words)
          - **summarised_text** (concise summary preserving legal meaning)

        - Maintain precise clause boundaries
        - Include all important clauses
        - Exclude non-contractual content
        - Response in **valid json** only

        Input: {text}

        Response format:
        [
          {{
            "clause_id": "<clause_id>",
            "heading/title": "<heading>",
            "summarised_text": "<summary>"
          }}
        ]
        """

        # Make API call with automatic retry and key rotation
        def api_call(client):
            chat_completion = client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.3-70b-versatile",
                response_format={"type": "json_object"},
                temperature=0.3
            )
            return chat_completion.choices[0].message.content
        
        return make_api_call_with_retry(api_call, max_retries=5)

    except Exception as e:
        raise Exception(f"Error during clause extraction: {str(e)}")

if __name__ == "__main__":
    try:
        pdf_file = "Data-Processing-Agreement-Template.pdf"
        json_file = "dpa.json"
        
        if extract_text_to_json(pdf_file, json_file):
            print(f"Successfully extracted text to {json_file}")
            
            # Test clause extraction
            result = Clause_extraction(pdf_file)
            print("Clause extraction successful!")
            print(result)
            
    except Exception as e:
        print(f"Error: {str(e)}")
